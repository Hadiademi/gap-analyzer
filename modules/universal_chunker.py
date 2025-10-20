"""
⚠️ BACKUP FILE - NOT CURRENTLY IN USE ⚠️

Alternative universal chunking implementation with document type detection.
Current production system uses: modules/gap_analyzer_claude.py

Kept for reference and testing purposes.
"""

from docx import Document as dx
from langchain.schema.document import Document
import re

def detect_document_type(doc):
    """
    Detect document type based on structure
    """
    has_list_paragraph = False
    has_headings = False
    has_control_pattern = False
    has_bold_titles = False
    
    for paragraph in doc.paragraphs[:50]:  # Check first 50 paragraphs
        text = paragraph.text.strip()
        style = paragraph.style.name
        
        if style == 'List Paragraph':
            has_list_paragraph = True
        
        if 'Heading' in style:
            has_headings = True
        
        if re.match(r'C-\d+\s*\|', text):  # Pattern like "C-01 |"
            has_control_pattern = True
        
        if any(run.bold for run in paragraph.runs) and text:
            has_bold_titles = True
    
    # Decision logic
    if has_control_pattern and has_headings:
        return "IT_RISK_CONTROLS"
    elif has_list_paragraph and has_bold_titles:
        return "CONCEPT_RISK"
    else:
        return "GENERIC"

def chunk_concept_risk(doc):
    """
    Chunking për Concept Risk and Governance document
    """
    content = []
    current_title = None
    current_subtitle = None
    current_data = []
    index = 1
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        style = paragraph.style.name
        has_bold = any(run.bold for run in paragraph.runs)
        
        # Main title (List Paragraph with number)
        if style == 'List Paragraph':
            # Save previous chunk
            if current_subtitle and current_data:
                content.append([current_title, current_subtitle, current_data])
                current_data = []
            
            current_title = str(index) + '. ' + text
            current_subtitle = None
            index += 1
            
        # Subtitle (Bold text)
        elif has_bold and style != 'List Paragraph':
            # Save previous chunk
            if current_subtitle and current_data:
                content.append([current_title, current_subtitle, current_data])
                current_data = []
            
            current_subtitle = text
            
        # Normal content
        elif current_title:
            if text:
                current_data.append(text)
    
    # Save last chunk
    if current_subtitle and current_data:
        content.append([current_title, current_subtitle, current_data])
    
    return content

def chunk_it_risk_controls(doc):
    """
    Chunking për IT Risk Controls document
    """
    content = []
    current_section = None
    current_control = None
    current_data = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        style = paragraph.style.name
        
        # Section header (Heading 2)
        if style == 'Heading 2':
            # Save previous chunk
            if current_control and current_data:
                content.append([current_section, current_control, current_data])
                current_data = []
            
            current_section = text
            current_control = None
            
        # Control (starts with C-XX | or List Bullet with C-XX)
        elif re.match(r'C-\d+\s*\|', text) or (style == 'List Bullet' and 'C-' in text):
            # Save previous chunk
            if current_control and current_data:
                content.append([current_section, current_control, current_data])
                current_data = []
            
            # Clean control text
            current_control = text.replace('**', '').strip()
            
        # Normal content (belongs to current control)
        elif current_control and style == 'Normal':
            if text:
                current_data.append(text)
    
    # Save last chunk
    if current_control and current_data:
        content.append([current_section, current_control, current_data])
    
    return content

def universal_chunk_document(file_path):
    """
    Universal chunking that detects document type and applies appropriate strategy
    """
    print(f"Processing: {file_path}")
    doc = dx(file_path)
    
    # Detect document type
    doc_type = detect_document_type(doc)
    print(f"Detected document type: {doc_type}")
    
    # Apply appropriate chunking
    if doc_type == "IT_RISK_CONTROLS":
        chunks = chunk_it_risk_controls(doc)
        print(f"Applied IT Risk Controls chunking")
    elif doc_type == "CONCEPT_RISK":
        chunks = chunk_concept_risk(doc)
        print(f"Applied Concept Risk chunking")
    else:
        # Fallback to generic chunking
        chunks = chunk_concept_risk(doc)  # Use as default
        print(f"Applied generic chunking")
    
    print(f"Total chunks created: {len(chunks)}")
    return chunks, doc_type

def create_embeddings_from_chunks(chunks, doc_type, embeddings):
    """
    Create embeddings for chunks based on document type
    """
    documents = []
    
    for i, chunk in enumerate(chunks):
        section_or_title = chunk[0] if chunk[0] else "General"
        control_or_subtitle = chunk[1] if chunk[1] else "No Subtitle"
        content_lines = chunk[2] if len(chunk) > 2 else []
        
        # Build content based on document type
        if doc_type == "IT_RISK_CONTROLS":
            # Format: Section → Control → Details
            joined_content = '\n'.join(content_lines)
            full_text = f"Section: {section_or_title}\nControl: {control_or_subtitle}\n{joined_content}"
            
            metadata = {
                'id': i + 1,
                'section': section_or_title,
                'control': control_or_subtitle,
                'type': 'it_control'
            }
        else:
            # Format: Title → Subtitle → Content
            joined_content = '\n'.join(content_lines)
            full_text = f"Title: {section_or_title}\nSubTitle: {control_or_subtitle}\n{joined_content}"
            
            metadata = {
                'id': i + 1,
                'title': section_or_title,
                'subtitle': control_or_subtitle,
                'type': 'concept_risk'
            }
        
        documents.append(Document(
            page_content=full_text,
            metadata=metadata
        ))
    
    return documents

def create_vectorstore_universal(file_path, db_directory, embeddings):
    """
    Create vector store using universal chunking
    """
    from langchain_chroma import Chroma
    import os
    
    # Chunk document
    chunks, doc_type = universal_chunk_document(file_path)
    
    # Create documents with embeddings
    documents = create_embeddings_from_chunks(chunks, doc_type, embeddings)
    
    # Create vector store
    if not os.path.isdir(db_directory):
        os.makedirs(db_directory)
    
    collection_name = f"universal_docs_{doc_type.lower()}"
    
    vectorstore = Chroma(
        collection_name=collection_name,
        embedding_function=embeddings,
        persist_directory=db_directory,
    )
    
    vectorstore.add_documents(documents)
    print(f"Vector store created with {len(documents)} documents")
    
    return vectorstore, doc_type

if __name__ == "__main__":
    # Test with both documents
    print("="*60)
    print("Testing IT Risk Controls Document")
    print("="*60)
    chunks1, type1 = universal_chunk_document("Data/Document/IT_Risk_Controls_for_Banks.docx")
    print(f"\nFirst 3 chunks:")
    for i, chunk in enumerate(chunks1[:3]):
        print(f"\nChunk {i+1}:")
        print(f"  Section/Title: {chunk[0]}")
        print(f"  Control/Subtitle: {chunk[1]}")
        print(f"  Content lines: {len(chunk[2])}")
    
    print("\n" + "="*60)
    print("Testing Concept Risk Document")
    print("="*60)
    chunks2, type2 = universal_chunk_document("Data/Document/Concept_Risk and Governance_EN.docx")
    print(f"\nFirst 3 chunks:")
    for i, chunk in enumerate(chunks2[:3]):
        print(f"\nChunk {i+1}:")
        print(f"  Section/Title: {chunk[0]}")
        print(f"  Control/Subtitle: {chunk[1]}")
        print(f"  Content lines: {len(chunk[2])}")
