"""
⚠️ BACKUP FILE - NOT CURRENTLY IN USE ⚠️

This file contains alternative chunking approaches with document type detection.
The current production system uses gap_analyzer_claude.py with universal chunking.

This file is kept for:
- Reference and comparison
- Testing alternative approaches
- Fallback if universal method has issues

Last used: Development phase
Current active file: modules/gap_analyzer_claude.py
"""

from docx import Document as dx
from langchain.schema.document import Document
import re
import os

def split_docx_by_bold_titles(file_path):
    """
    Original chunking për Concept Risk document (nga analyzer.py)
    Funksionon për dokumente me:
    - List Paragraph (1., 2., 3.)
    - Bold text për subtitles
    - Normal text për content
    """
    doc = dx(file_path)
    
    content = []
    current_title = None
    index = 1
    main_title = None

    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'List Paragraph':
            main_title = str(index) + '. ' + paragraph.text.strip()
            index += 1
            content.append([main_title, "", []])
            current_title = None
            
        elif paragraph.style.name != 'List Paragraph' and any(run.bold for run in paragraph.runs):
            current_title = paragraph.text.strip()
            content.append([main_title, current_title, []])
            
        elif main_title:
            text = paragraph.text.strip()
            if text:
                if current_title:
                    content[-1][2].append(text)
                else:
                    content[-1][2].append(text)

    return content

def split_it_risk_controls(file_path):
    """
    Chunking për IT Risk Controls document
    Funksionon për dokumente me:
    - Heading 2 për seksione
    - List Bullet me C-XX pattern për kontrolle
    - Normal text për detaje
    """
    doc = dx(file_path)
    
    content = []
    current_section = None
    current_control = None
    current_data = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        style = paragraph.style.name
        
        # Section header (Heading 2)
        if style == 'Heading 2':
            # Save previous chunk
            if current_control and current_data:
                content.append([current_section, current_control, current_data])
                current_data = []
            
            current_section = text
            current_control = None
            
        # Control (C-XX pattern)
        elif re.match(r'C-\d+\s*\|', text) or (style == 'List Bullet' and 'C-' in text):
            # Save previous chunk
            if current_control and current_data:
                content.append([current_section, current_control, current_data])
                current_data = []
            
            current_control = text.replace('**', '').strip()
            
        # Normal content
        elif current_control and style == 'Normal':
            if text:
                current_data.append(text)
    
    # Save last chunk
    if current_control and current_data:
        content.append([current_section, current_control, current_data])
    
    return content

def detect_document_type(file_path):
    """
    Detect nëse është Concept Risk apo IT Risk Controls
    """
    doc = dx(file_path)
    
    has_list_paragraph = False
    has_heading2 = False
    has_control_pattern = False
    
    for paragraph in doc.paragraphs[:30]:  # Check first 30 paragraphs
        text = paragraph.text.strip()
        style = paragraph.style.name
        
        if style == 'List Paragraph':
            has_list_paragraph = True
        
        if style == 'Heading 2':
            has_heading2 = True
        
        if re.match(r'C-\d+\s*\|', text):
            has_control_pattern = True
    
    # Decision
    if has_control_pattern and has_heading2:
        return "IT_RISK"
    elif has_list_paragraph:
        return "CONCEPT_RISK"
    else:
        return "CONCEPT_RISK"  # Default fallback

def universal_split_document(file_path):
    """
    Universal function që zgjedh metodën e duhur
    """
    doc_type = detect_document_type(file_path)
    print(f"Detected: {doc_type}")
    
    if doc_type == "IT_RISK":
        chunks = split_it_risk_controls(file_path)
    else:
        chunks = split_docx_by_bold_titles(file_path)
    
    print(f"Created {len(chunks)} chunks")
    return chunks, doc_type

def create_vectorstore(db_directory, split_content, embeddings, doc_type="CONCEPT_RISK"):
    """
    Create vector store (nga analyzer.py, përmirësuar)
    """
    from langchain_chroma import Chroma
    
    if not os.path.isdir(db_directory):
        os.makedirs(db_directory)
    
    to_embed_docs = []
    doc_num = 1
    
    for row in split_content:
        if row[2]:  # If has content
            joined_par = '\n'.join(row[2])
            
            # Build content based on document type
            if doc_type == "IT_RISK":
                # Format: Section → Control → Details
                section = row[0] if row[0] else "General"
                control = row[1] if row[1] else "No Control"
                content = f"Section: {section}\nControl: {control}\n{joined_par}"
                
                metadata = {
                    'id': doc_num,
                    'section': section,
                    'control': control,
                    'type': 'it_risk'
                }
            else:
                # Format: Title → SubTitle → Content (original)
                title = row[0] if row[0] else "No Title"
                subtitle = row[1] if row[1] else ""
                
                if subtitle:
                    content = f"Title: {title}\n SubTitle: {subtitle}\n{joined_par}"
                else:
                    content = f"Title: {title}\n{joined_par}"
                
                metadata = {
                    'id': doc_num,
                    'title': title,
                    'sub_title': subtitle,
                    'type': 'concept_risk'
                }
            
            to_embed_docs.append(Document(
                page_content=content,
                metadata=metadata
            ))
            doc_num += 1
    
    # Create ChromaDB
    vectorstore = Chroma(
        collection_name="collection_document",
        embedding_function=embeddings,
        persist_directory=db_directory,
    )
    vectorstore.add_documents(to_embed_docs)
    
    return vectorstore

# Test function
if __name__ == "__main__":
    print("="*60)
    print("TEST 1: Concept Risk Document")
    print("="*60)
    
    chunks1, type1 = universal_split_document("Data/Document/Concept_Risk and Governance_EN.docx")
    print(f"Type: {type1}")
    print(f"Total chunks: {len(chunks1)}")
    
    if len(chunks1) > 0:
        print("\nFirst 3 chunks:")
        for i in range(min(3, len(chunks1))):
            print(f"\nChunk {i+1}:")
            print(f"  Title: {chunks1[i][0]}")
            print(f"  Subtitle: {chunks1[i][1]}")
            print(f"  Content lines: {len(chunks1[i][2])}")
            if chunks1[i][2]:
                print(f"  First line: {chunks1[i][2][0][:80]}...")
    
    print("\n" + "="*60)
    print("TEST 2: IT Risk Controls Document")
    print("="*60)
    
    chunks2, type2 = universal_split_document("Data/Document/IT_Risk_Controls_for_Banks.docx")
    print(f"Type: {type2}")
    print(f"Total chunks: {len(chunks2)}")
    
    if len(chunks2) > 0:
        print("\nFirst 3 chunks:")
        for i in range(min(3, len(chunks2))):
            print(f"\nChunk {i+1}:")
            print(f"  Section: {chunks2[i][0]}")
            print(f"  Control: {chunks2[i][1]}")
            print(f"  Content lines: {len(chunks2[i][2])}")
            if chunks2[i][2]:
                print(f"  First line: {chunks2[i][2][0][:80]}...")
